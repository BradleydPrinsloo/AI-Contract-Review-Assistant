from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from .branding import DECISION_SUPPORT_NOTICE, REPORT_FULL_TITLE


@dataclass
class ScanResult:
    phrase: str
    category: str
    finding_type: str
    risk: str
    score: int
    confidence: int
    location: str
    note: str
    context: str
    reason: str = ""


def load_keywords(path: str | Path) -> list[dict]:
    payload = json.loads(Path(path).read_text(encoding="utf-8"))
    if not isinstance(payload, list):
        raise ValueError("Keyword library must contain a JSON array.")
    return payload


def extract_document(path: str | Path) -> list[tuple[str, str]]:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".txt":
        return [("Document", source.read_text(encoding="utf-8", errors="ignore"))]
    if suffix == ".docx":
        from docx import Document
        document = Document(str(source))
        return [(f"Paragraph {index}", paragraph.text) for index, paragraph in enumerate(document.paragraphs, 1) if paragraph.text.strip()]
    if suffix == ".pdf":
        return _extract_pdf(source)
    raise ValueError("Supported formats are PDF, DOCX, and TXT.")


def _extract_pdf(path: Path) -> list[tuple[str, str]]:
    import fitz
    document = fitz.open(path)
    chunks: list[tuple[str, str]] = []
    for page_index, page in enumerate(document, 1):
        text = page.get_text("text").strip()
        if not text:
            text = _ocr_page(page)
        if text:
            chunks.append((f"Page {page_index}", text))
    return chunks


def _ocr_page(page) -> str:
    try:
        from rapidocr_onnxruntime import RapidOCR
        pix = page.get_pixmap(matrix=__import__("fitz").Matrix(2, 2), alpha=False)
        result, _ = RapidOCR()(pix.tobytes("png"))
        return "\n".join(item[1] for item in (result or []) if len(item) > 1)
    except Exception:
        return ""


def scan_chunks(chunks: Iterable[tuple[str, str]], rules: Iterable[dict]) -> list[ScanResult]:
    findings: list[ScanResult] = []
    seen: set[tuple[str, str, str]] = set()
    for location, text in chunks:
        for sentence in _sentences(text):
            for rule in rules:
                phrase = str(rule.get("phrase", "")).strip()
                aliases = [str(value).strip() for value in rule.get("aliases", []) if str(value).strip()]
                terms = [phrase, *aliases]
                matched = next((term for term in terms if term and re.search(rf"\b{re.escape(term)}\b", sentence, re.I)), None)
                if not matched:
                    continue
                finding_type = str(rule.get("finding_type", "Risk"))
                if _is_nonoperative_reference(sentence, matched, finding_type):
                    continue
                key = (location, phrase.casefold(), sentence.casefold())
                if key in seen:
                    continue
                seen.add(key)
                risk = str(rule.get("risk", "Medium"))
                confidence = _confidence(sentence, matched, finding_type)
                findings.append(ScanResult(
                    phrase or matched,
                    str(rule.get("category", "General")),
                    finding_type,
                    risk,
                    0,
                    confidence,
                    location,
                    str(rule.get("note", "Review the surrounding clause manually.")),
                    sentence.strip(),
                    f"Matched configured rule: {matched}",
                ))
    return findings


def _sentences(text: str) -> list[str]:
    compact = re.sub(r"\s+", " ", text).strip()
    return [part.strip() for part in re.split(r"(?<=[.!?;])\s+", compact) if part.strip()]


def _is_nonoperative_reference(sentence: str, matched: str, finding_type: str) -> bool:
    lowered = sentence.casefold()
    operative = any(token in lowered for token in (
        " shall ", " must ", " agrees to ", " will ", " is required to ",
        " may terminate", " is liable", " shall be liable", " warrants ",
        " represents ", " indemnifies ", " shall indemnify", " agrees that",
    ))
    casual_markers = (
        "for discussion", "discussion only", "mentioned", "reference only",
        "example", "illustration", "generally", "may include", "could include",
        "does not create", "does not establish", "is not intended to",
        "no party is obligated", "no obligation", "not binding", "draft note",
        "training material", "background information", "considered but rejected",
    )
    negated_obligation = any(token in lowered for token in (
        "shall not", "not required", "does not apply", "is not liable",
        "will not be liable", "no duty to", "no obligation to",
    ))
    if negated_obligation and finding_type.casefold() == "risk":
        return True
    if any(marker in lowered for marker in casual_markers) and not operative:
        return True
    if not operative and len(sentence.split()) < 8 and finding_type.casefold() == "risk":
        return True
    return False


def _confidence(sentence: str, matched: str, finding_type: str) -> int:
    lowered = sentence.lower()
    score = 60
    if any(word in lowered for word in ("shall", "must", "agrees", "required", "may terminate", "liable", "warrants")):
        score += 25
    if any(word in lowered for word in ("not required", "shall not", "does not apply", "for discussion only", "does not create")):
        score -= 30
    if finding_type == "Info":
        score = min(score, 55)
    if len(matched.split()) >= 3:
        score += 5
    return max(10, min(score, 98))


def export_csv(results: Iterable[ScanResult], output_path: str | Path, source_file: str = "", risk_assessment=None) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Source", source_file])
        if risk_assessment is not None:
            writer.writerow(["Risk Score", risk_assessment.total_score, "Rating", risk_assessment.rating])
        writer.writerow([])
        writer.writerow(["Phrase", "Category", "Finding Type", "Risk", "Score", "Confidence", "Location", "Note", "Context"])
        for item in results:
            writer.writerow([item.phrase, item.category, item.finding_type, item.risk, item.score, item.confidence, item.location, item.note, item.context])
    return output


def export_txt(results: Iterable[ScanResult], source_file: str, output_path: str | Path, risk_assessment, summary_text: str = "") -> Path:
    output = Path(output_path)
    lines = [REPORT_FULL_TITLE, "========================", f"Source file: {source_file}", f"Generated: {datetime.now().isoformat(timespec='seconds')}", f"Overall Risk Score: {risk_assessment.total_score}/100", f"Risk Rating: {risk_assessment.rating}", f"Finding Count: {risk_assessment.finding_count}", f"Risk Findings: {risk_assessment.risk_count}", f"Protective Findings: {risk_assessment.protective_count}", f"Neutral/Info Findings: {risk_assessment.neutral_count}", "", summary_text.strip(), "", "Detailed findings", "-----------------"]
    for index, item in enumerate(results, 1):
        lines.extend([f"{index}. {item.phrase} [{item.finding_type}/{item.risk}]", f"Category: {item.category}", f"Location: {item.location}", f"Score: {item.score} | Confidence: {item.confidence}%", f"Note: {item.note}", f"Context: {item.context}", ""])
    lines.append(f"Disclaimer: {DECISION_SUPPORT_NOTICE}")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines), encoding="utf-8")
    return output


def export_docx(results: Iterable[ScanResult], source_file: str, output_path: str | Path, risk_assessment, summary_text: str = "") -> Path:
    from docx import Document
    from docx.shared import Pt
    output = Path(output_path)
    document = Document()
    document.add_heading(REPORT_FULL_TITLE, 0)
    document.add_paragraph(f"Source: {source_file}")
    document.add_paragraph(f"Generated: {datetime.now().isoformat(timespec='seconds')}")
    document.add_heading("Risk Overview", level=1)
    document.add_paragraph(f"Score: {risk_assessment.total_score}/100\nRating: {risk_assessment.rating}\nFindings: {risk_assessment.finding_count}")
    if summary_text.strip():
        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(summary_text.strip())
    document.add_heading("Detailed Findings", level=1)
    for item in results:
        document.add_heading(f"{item.phrase} — {item.category}", level=2)
        document.add_paragraph(f"Type: {item.finding_type} | Risk: {item.risk} | Score: {item.score} | Confidence: {item.confidence}%")
        document.add_paragraph(f"Location: {item.location}")
        document.add_paragraph(f"Review note: {item.note}")
        document.add_paragraph(item.context)
    disclaimer = document.add_paragraph(DECISION_SUPPORT_NOTICE)
    for run in disclaimer.runs:
        run.font.size = Pt(9)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
