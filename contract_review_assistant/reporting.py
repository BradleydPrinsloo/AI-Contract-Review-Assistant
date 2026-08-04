from __future__ import annotations

import html
from datetime import datetime
from pathlib import Path
from typing import Iterable

from contract_review_assistant.branding import (
    DECISION_SUPPORT_NOTICE,
    PRODUCT_NAME,
    REPORT_EXECUTIVE_TITLE,
    REPORT_FULL_TITLE,
)
from contract_review_assistant.scanner import ScanResult


def build_report_html(
    results: Iterable[ScanResult],
    source_file: str,
    risk_assessment,
    summary_text: str,
    *,
    report_type: str = "full",
    include_summary: bool = True,
    include_findings: bool = True,
    include_recommendations: bool = True,
) -> str:
    """Build printable HTML for PDF generation and report previews."""

    findings = list(results)
    title = (
        REPORT_EXECUTIVE_TITLE
        if report_type == "executive"
        else REPORT_FULL_TITLE
    )
    generated = datetime.now().strftime("%B %d, %Y at %I:%M %p")
    source_name = Path(source_file).name if source_file else "Unknown contract"

    high = sum(
        1
        for item in findings
        if item.risk.casefold() in {"critical", "high", "elevated"}
    )
    medium = sum(
        1 for item in findings if item.risk.casefold() in {"moderate", "medium"}
    )
    low = len(findings) - high - medium

    sections = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        "<style>",
        "body{font-family:Segoe UI,Arial,sans-serif;color:#172033;margin:34px;}",
        "h1{font-size:26px;margin-bottom:4px;color:#0f172a;}",
        "h2{font-size:18px;margin-top:26px;border-bottom:1px solid #cbd5e1;padding-bottom:6px;}",
        ".brand{color:#0369a1;font-size:11px;font-weight:900;letter-spacing:2px;text-transform:uppercase;margin-bottom:8px;}",
        ".muted{color:#64748b;font-size:11px;}",
        ".cards{width:100%;border-collapse:separate;border-spacing:8px;margin:18px 0;}",
        ".card{border:1px solid #cbd5e1;border-radius:8px;padding:12px;background:#f8fafc;}",
        ".label{font-size:10px;color:#64748b;text-transform:uppercase;font-weight:700;}",
        ".value{font-size:20px;font-weight:800;margin-top:4px;}",
        "table.findings{width:100%;border-collapse:collapse;font-size:10px;}",
        "table.findings th{background:#e2e8f0;text-align:left;padding:7px;border:1px solid #cbd5e1;}",
        "table.findings td{padding:7px;vertical-align:top;border:1px solid #cbd5e1;}",
        ".high{color:#c2410c;font-weight:700}.medium{color:#a16207;font-weight:700}.low{color:#15803d;font-weight:700}",
        ".disclaimer{margin-top:28px;padding:10px;border-left:4px solid #94a3b8;background:#f8fafc;font-size:9px;color:#475569;}",
        "</style></head><body>",
        f"<div class='brand'>{html.escape(PRODUCT_NAME)}</div>",
        f"<h1>{html.escape(title)}</h1>",
        f"<div class='muted'>Contract: {html.escape(source_name)}<br>Generated: {html.escape(generated)}</div>",
        "<table class='cards'><tr>",
        f"<td class='card'><div class='label'>Overall Risk</div><div class='value'>{risk_assessment.total_score}/100</div></td>",
        f"<td class='card'><div class='label'>Rating</div><div class='value'>{html.escape(str(risk_assessment.rating))}</div></td>",
        f"<td class='card'><div class='label'>Total Findings</div><div class='value'>{len(findings)}</div></td>",
        f"<td class='card'><div class='label'>Severity</div><div class='value'>{high} H / {medium} M / {low} L</div></td>",
        "</tr></table>",
    ]

    if include_summary:
        sections.extend(
            [
                "<h2>Executive Summary</h2>",
                _text_to_html(summary_text or "No executive summary is available."),
            ]
        )

    if report_type == "executive":
        priorities = sorted(findings, key=lambda item: item.score, reverse=True)[:5]
        sections.append("<h2>Priority Review Items</h2>")
        if priorities:
            sections.append("<ol>")
            for item in priorities:
                sections.append(
                    "<li><b>{}</b> — {} / {} (score {})<br>{}</li>".format(
                        html.escape(item.phrase),
                        html.escape(item.category),
                        html.escape(item.risk),
                        item.score,
                        html.escape(item.note),
                    )
                )
            sections.append("</ol>")
        else:
            sections.append("<p>No priority findings were detected.</p>")
    elif include_findings:
        sections.append("<h2>Detailed Findings</h2>")
        sections.append(
            "<table class='findings'><tr><th>Clause</th><th>Category</th><th>Risk</th><th>Score</th><th>Location</th><th>Review Guidance</th><th>Context</th></tr>"
        )
        for item in findings:
            risk_class = _risk_class(item.risk)
            sections.append(
                "<tr><td>{}</td><td>{}</td><td class='{}'>{}</td><td>{}</td><td>{}</td><td>{}</td><td>{}</td></tr>".format(
                    html.escape(item.phrase),
                    html.escape(item.category),
                    risk_class,
                    html.escape(item.risk),
                    item.score,
                    html.escape(item.location),
                    html.escape(item.note),
                    html.escape(item.context),
                )
            )
        sections.append("</table>")

    if include_recommendations:
        sections.extend(
            [
                "<h2>Recommended Next Step</h2>",
                "<p>Review all high-priority clauses, compare the language against the organization's approved playbook, and obtain qualified legal review before approval or signature.</p>",
            ]
        )

    sections.extend(
        [
            f"<div class='disclaimer'><b>Decision-support notice:</b> {html.escape(DECISION_SUPPORT_NOTICE)}</div>",
            "</body></html>",
        ]
    )
    return "".join(sections)


def export_report_docx(
    results: Iterable[ScanResult],
    source_file: str,
    output_path: str | Path,
    risk_assessment,
    summary_text: str,
    *,
    report_type: str = "full",
    include_summary: bool = True,
    include_findings: bool = True,
    include_recommendations: bool = True,
) -> Path:
    """Generate an executive or full professional Word report."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    findings = list(results)
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)

    title = (
        REPORT_EXECUTIVE_TITLE
        if report_type == "executive"
        else REPORT_FULL_TITLE
    )
    heading = document.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_line = document.add_paragraph(PRODUCT_NAME)
    brand_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"Contract: {Path(source_file).name if source_file else 'Unknown contract'}")
    document.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

    overview = document.add_table(rows=2, cols=4)
    overview.style = "Table Grid"
    labels = ["Overall Risk", "Rating", "Findings", "High Priority"]
    high = sum(
        1
        for item in findings
        if item.risk.casefold() in {"critical", "high", "elevated"}
    )
    values = [
        f"{risk_assessment.total_score}/100",
        str(risk_assessment.rating),
        str(len(findings)),
        str(high),
    ]
    for col, label in enumerate(labels):
        overview.cell(0, col).text = label
        overview.cell(1, col).text = values[col]

    if include_summary:
        document.add_heading("Executive Summary", level=1)
        document.add_paragraph(summary_text or "No executive summary is available.")

    if report_type == "executive":
        document.add_heading("Priority Review Items", level=1)
        priorities = sorted(findings, key=lambda item: item.score, reverse=True)[:5]
        if priorities:
            for item in priorities:
                document.add_paragraph(
                    f"{item.phrase} — {item.category} / {item.risk} (score {item.score})\n{item.note}",
                    style="List Number",
                )
        else:
            document.add_paragraph("No priority findings were detected.")
    elif include_findings:
        document.add_heading("Detailed Findings", level=1)
        for item in findings:
            document.add_heading(f"{item.phrase} — {item.category}", level=2)
            document.add_paragraph(
                f"Type: {item.finding_type} | Risk: {item.risk} | Score: {item.score} | Confidence: {item.confidence}%"
            )
            document.add_paragraph(f"Location: {item.location}")
            document.add_paragraph(f"Review guidance: {item.note}")
            document.add_paragraph(f"Context: {item.context}")

    if include_recommendations:
        document.add_heading("Recommended Next Step", level=1)
        document.add_paragraph(
            "Review all high-priority clauses, compare the language against the organization's approved playbook, and obtain qualified legal review before approval or signature."
        )

    disclaimer = document.add_paragraph(
        f"Decision-support notice: {DECISION_SUPPORT_NOTICE}"
    )
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.italic = True

    document.save(output)
    return output


def _text_to_html(text: str) -> str:
    paragraphs: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or set(stripped) == {"="}:
            continue
        if stripped.startswith("-"):
            paragraphs.append(f"<li>{html.escape(stripped[1:].strip())}</li>")
        else:
            paragraphs.append(f"<p>{html.escape(stripped)}</p>")
    return "".join(paragraphs)


def _risk_class(risk: str) -> str:
    lowered = risk.casefold()
    if lowered in {"critical", "high", "elevated"}:
        return "high"
    if lowered in {"moderate", "medium"}:
        return "medium"
    return "low"
